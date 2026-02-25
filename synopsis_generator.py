from docx import Document

class SynopsisGenerator:
    def __init__(self, template_path):
        self.template_path = template_path

    def fill_template(self, data):
        doc = Document(self.template_path)

        # Замена в параграфах
        for para in doc.paragraphs:
            self._replace_in_para(para, data)

        # Замена в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_para(para, data)

        return doc

    def _replace_in_para(self, para, data):
        for key, value in data.items():
            placeholder = f'{{{{{key}}}}}'
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, str(value))

    def save_docx(self, doc, filename):
        doc.save(filename)